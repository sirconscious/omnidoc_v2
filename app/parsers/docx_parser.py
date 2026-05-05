import logging
from pathlib import Path
from typing import Tuple, List, Dict, Any
from datetime import datetime

logger = logging.getLogger(__name__)
MIN_WORDS_PER_CHUNK = 20

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    Document = None
    HAS_DOCX = False
    logger.warning("python-docx not installed - DOCX parsing unavailable")


def _extract_headers_footers(doc) -> Tuple[List[str], List[str]]:
    """Extract headers and footers from each section."""
    headers = []
    footers = []
    
    try:
        if hasattr(doc, 'sections'):
            for section in doc.sections:
                try:
                    header = section.header
                    if header:
                        header_texts = [p.text for p in header.paragraphs if p.text.strip()]
                        if header_texts:
                            headers.extend(header_texts)
                except Exception as e:
                    logger.debug(f"Could not extract header: {e}")
                
                try:
                    footer = section.footer
                    if footer:
                        footer_texts = [p.text for p in footer.paragraphs if p.text.strip()]
                        if footer_texts:
                            footers.extend(footer_texts)
                except Exception as e:
                    logger.debug(f"Could not extract footer: {e}")
    except Exception as e:
        logger.warning(f"Failed to extract headers/footers: {e}")
    
    return headers, footers


def _extract_text_boxes(doc) -> List[str]:
    """Extract text from text boxes/shapes in the document."""
    text_boxes = []
    
    try:
        body = doc.element.body
        for elem in body.iter():
            if elem.tag.endswith('}txbx'):
                for text in elem.iter():
                    if text.tag.endswith('}t') and text.text:
                        text_boxes.append(text.text)
    except Exception as e:
        logger.debug(f"Could not extract text boxes: {e}")
    
    return text_boxes


def _format_table_markdown(table) -> str:
    """Format a table as markdown."""
    if not table.rows:
        return ""
    
    rows_data = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows_data.append(cells)
    
    if not rows_data:
        return ""
    
    col_count = len(rows_data[0])
    col_widths = []
    for col_idx in range(col_count):
        max_w = len(rows_data[0][col_idx])
        for row in rows_data:
            if col_idx < len(row):
                max_w = max(max_w, len(row[col_idx]))
        col_widths.append(max_w)
    
    lines = []
    
    lines.append("|" + "|".join([str(rows_data[0][i]).center(col_widths[i]) for i in range(col_count)]) + "|")
    lines.append("|" + "|".join(["-" * (w + 2) for w in col_widths]) + "|")
    
    for row in rows_data[1:]:
        lines.append("|" + "|".join([str(row[i]).ljust(col_widths[i]) if i < len(row) else " " * col_widths[i] for i in range(col_count)]) + "|")
    
    return "\n".join(lines)


def _extract_core_properties(doc) -> Dict[str, Any]:
    """Extract core properties (author, created, modified) from docx."""
    props = {
        "author": None,
        "created": None,
        "modified": None,
        "last_modified_by": None,
        "title": None,
    }
    
    try:
        core_props = doc.core_properties
        props["author"] = getattr(core_props, 'author', None) or None
        props["created"] = getattr(core_props, 'created', None)
        props["modified"] = getattr(core_props, 'modified', None)
        props["last_modified_by"] = getattr(core_props, 'last_modified_by', None) or None
        props["title"] = getattr(core_props, 'title', None) or None
        
        for key in ("created", "modified"):
            if props[key]:
                try:
                    props[key] = str(props[key])
                    has_separator = "T" in props[key] or " " in props[key]
                    if has_separator:
                        separator = "T" if "T" in props[key] else " "
                        year_str = props[key].split(separator)[0].split("-")[0]
                        year = int(year_str)
                        if year < 2015:
                            props[key] = None
                except Exception:
                    props[key] = None
        
        for key in ("author", "last_modified_by", "title"):
            if props[key]:
                props[key] = str(props[key])
    except Exception as e:
        logger.debug(f"Could not extract core properties: {e}")
    
    return props


def _has_table_in_text(text: str) -> bool:
    """Check if chunk text contains table markup."""
    return "--- Table" in text or "| " in text[:200]


def _is_resume_section_header(text: str) -> bool:
    """Detect if a paragraph is a resume section header."""
    if not text or len(text.split()) > 6:
        return False

    text_lower = text.lower().strip()

    resume_headers = {
        'professional experience', 'work experience', 'experience',
        'employment history', 'professional history',
        'education', 'academic background', 'qualifications',
        'skills', 'technical skills', 'core competencies', 'expertise',
        'summary', 'professional summary', 'profile', 'objective',
        'projects', 'key projects', 'notable projects',
        'certifications', 'certificates', 'licenses',
        'achievements', 'accomplishments', 'awards',
        'publications', 'research', 'patents',
        'languages', 'professional affiliations', 'memberships',
    }

    for header in resume_headers:
        if text_lower == header or text_lower.startswith(header):
            return True

    return False


def _chunk_resume_by_sections(doc) -> List[Dict[str, Any]]:
    """Chunk resume by semantic sections instead of word count."""
    chunks = []
    current_section_lines = []
    current_section_heading = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""

        is_header = (
            style_name.startswith('Heading')
            or _is_resume_section_header(text)
        )

        if is_header:
            if current_section_lines:
                section_text = "\n".join(current_section_lines)
                word_count = len(section_text.split())

                if word_count >= 20:
                    chunks.append({
                        "index": len(chunks),
                        "text": section_text,
                        "source_page": None,
                        "source_section": current_section_heading,
                        "has_table": _has_table_in_text(section_text),
                        "word_count": word_count,
                    })

            current_section_heading = text
            current_section_lines = [text]
        else:
            current_section_lines.append(text)

    if current_section_lines:
        section_text = "\n".join(current_section_lines)
        word_count = len(section_text.split())

        if word_count >= 20:
            chunks.append({
                "index": len(chunks),
                "text": section_text,
                "source_page": None,
                "source_section": current_section_heading,
                "has_table": _has_table_in_text(section_text),
                "word_count": word_count,
            })

    return chunks


def _is_likely_resume(doc) -> bool:
    """Heuristic to detect if a DOCX file is likely a resume/CV."""
    full_text = "\n".join([p.text for p in doc.paragraphs[:20]]).lower()

    resume_indicators = [
        'experience', 'education', 'skills', 'resume',
        'curriculum vitae', 'cv', 'objective', 'summary',
        'email:', 'phone:', 'linkedin',
    ]

    indicator_count = sum(1 for indicator in resume_indicators if indicator in full_text)
    return indicator_count >= 3


def _extract_sections(doc) -> List[Dict[str, Any]]:
    """Extract document sections/headings for chunking."""
    sections = []
    current_section = {"heading": None, "paragraphs": []}
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        style_name = para.style.name if para.style else ""
        if style_name.startswith('Heading'):
            if current_section["paragraphs"]:
                sections.append(current_section)
            current_section = {"heading": text, "paragraphs": [text]}
        else:
            if current_section["heading"] is None and current_section["paragraphs"]:
                current_section["paragraphs"].append(text)
            elif current_section["heading"] is None:
                current_section["paragraphs"] = [text]
            else:
                current_section["paragraphs"].append(text)
    
    if current_section["paragraphs"]:
        sections.append(current_section)
    
    return sections


def _filter_small_chunks(chunks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Filter out chunks with less than MIN_WORDS_PER_CHUNK words."""
    filtered = [c for c in chunks if c.get("word_count", 0) >= MIN_WORDS_PER_CHUNK]
    
    if not filtered and chunks:
        max_chunk = max(chunks, key=lambda c: c.get("word_count", 0))
        filtered = [max_chunk]
    
    for i, chunk in enumerate(filtered):
        chunk["index"] = i
    
    return filtered


def parse_docx(file_path: Path) -> Tuple[str, Dict[str, Any], List[Dict[str, Any]]]:
    """Parse a DOCX file with enhanced extraction."""
    if not HAS_DOCX:
        logger.error("python-docx not available")
        return "", {"error": "python-docx not installed"}, []
    
    doc = Document(file_path)
    
    metadata = {
        "paragraphs": 0,
        "tables": 0,
        "has_headers": False,
        "has_footers": False,
        "has_text_boxes": False,
        "author": None,
        "created": None,
        "modified": None,
        "last_modified_by": None,
        "title": None,
    }
    
    text_parts = []
    chunks = []
    
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text_parts.extend(paragraphs)
    metadata["paragraphs"] = len(paragraphs)
    
    core_props = _extract_core_properties(doc)
    metadata.update(core_props)
    
    headers, footers = _extract_headers_footers(doc)
    if headers:
        metadata["has_headers"] = True
        text_parts.insert(0, "--- HEADER ---\n" + "\n".join(headers))
    if footers:
        metadata["has_footers"] = True
        text_parts.append("--- FOOTER ---\n" + "\n".join(footers))
    
    text_boxes = _extract_text_boxes(doc)
    if text_boxes:
        metadata["has_text_boxes"] = True
        text_parts.extend(text_boxes)
    
    table_texts = []
    for table in doc.tables:
        metadata["tables"] += 1
        table_md = _format_table_markdown(table)
        if table_md:
            table_texts.append(table_md)
    
    # Fix contradiction: has_tables derived from tables count
    metadata["has_tables"] = metadata["tables"] > 0
    
    is_resume = _is_likely_resume(doc)

    if is_resume:
        logger.info("Detected resume format - using section-based chunking")
        chunks = _chunk_resume_by_sections(doc)
    else:
        sections = _extract_sections(doc)

        for idx, section in enumerate(sections):
            section_text = "\n".join(section["paragraphs"])
            if section_text:
                chunk_has_table = _has_table_in_text(section_text)
                chunks.append({
                    "index": idx,
                    "text": section_text,
                    "source_page": None,
                    "source_section": section["heading"],
                    "has_table": chunk_has_table,
                    "word_count": len(section_text.split()),
                })
    
    for table_md in table_texts:
        chunks.append({
            "index": len(chunks),
            "text": table_md,
            "source_page": None,
            "source_section": "table",
            "has_table": True,
            "word_count": len(table_md.split()),
        })
    
    if not chunks and paragraphs:
        all_text = "\n".join(paragraphs)
        chunks.append({
            "index": 0,
            "text": all_text,
            "source_page": None,
            "source_section": None,
            "has_table": False,
            "word_count": len(all_text.split()),
        })
    
    # Note: filtering now handled in main.py after overlap
    # chunks = _filter_small_chunks(chunks)
    
    full_text = "\n".join(text_parts)
    
    if table_texts:
        full_text += "\n\n" + "\n\n".join(table_texts)
    
    return full_text, metadata, chunks
